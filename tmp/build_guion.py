from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "Guion_exposicion_MLflow_Jedi_15-20_min.docx"

NAVY = RGBColor(21, 34, 56)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GOLD = RGBColor(166, 124, 0)
PALE_GOLD = "F5EFD7"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = RGBColor(92, 99, 112)
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(20, 20, 20)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_rich_paragraph(doc, text, style=None, before=0, after=6, line=1.25, keep=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_together = keep
    chunks = text.split("**")
    for i, chunk in enumerate(chunks):
        if not chunk:
            continue
        r = p.add_run(chunk)
        set_run_font(r, size=11, color=BLACK, bold=(i % 2 == 1))
    return p


def add_cue(doc, text, label="EN PANTALLA"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360], indent_dxa=120)
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(f"{label}  ")
    set_run_font(r, size=9.5, color=DARK_BLUE, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=9.5, color=BLACK, italic=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_anchor(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360], indent_dxa=120)
    cell = table.cell(0, 0)
    shade_cell(cell, PALE_GOLD)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run("FRASE ANCLA  ")
    set_run_font(r, size=9.5, color=GOLD, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10, color=NAVY, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_section_title(doc, slide, title, timing):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"Diapositiva {slide} · {title}")
    set_run_font(r, size=16, color=BLUE, bold=True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    r = p2.add_run(f"Tiempo orientativo: {timing}")
    set_run_font(r, size=9.5, color=MUTED, italic=True)


def add_bullet(doc, text):
    p = add_rich_paragraph(doc, text, style="List Bullet", after=4, line=1.25)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    return p


def add_number(doc, text):
    p = add_rich_paragraph(doc, text, style="List Number", after=4, line=1.25)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = True
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for list_style in ("List Bullet", "List Number"):
    s = styles[list_style]
    s.font.name = "Calibri"
    s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    s.font.size = Pt(11)
    s.paragraph_format.space_after = Pt(4)
    s.paragraph_format.line_spacing = 1.25

header = section.header
hp = header.paragraphs[0]
hp.text = "GUION PERSONAL · MLFLOW JEDI"
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_run_font(hp.runs[0], size=8.5, color=MUTED, bold=True)
footer = section.footer
add_page_field(footer.paragraphs[0])

# Opening block: workshop_agenda pattern with a restrained Star Wars color override.
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(22)
p.paragraph_format.space_after = Pt(2)
r = p.add_run("GUION DE EXPOSICIÓN")
set_run_font(r, size=11, color=GOLD, bold=True)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(5)
r = p.add_run("MLflow Jedi")
set_run_font(r, size=30, color=NAVY, bold=True)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
r = p.add_run("La Cámara de los Experimentos Perdidos")
set_run_font(r, size=18, color=DARK_BLUE, bold=False)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(18)
r = p.add_run("Un guion oral, humanizado y completo para acompañar la presentación y la demostración de la app.")
set_run_font(r, size=11.5, color=MUTED, italic=True)

timing = doc.add_table(rows=2, cols=3)
timing.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(timing, [3120, 3120, 3120], indent_dxa=120)
for idx, (label, value) in enumerate((("MLFLOW", "≈ 8 min"), ("APP", "≈ 8 min"), ("CIERRE", "≈ 1,5 min"))):
    c1 = timing.cell(0, idx)
    c2 = timing.cell(1, idx)
    shade_cell(c1, "152238")
    shade_cell(c2, PALE_GOLD)
    for cell in (c1, c2):
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c1.paragraphs[0].add_run(label)
    set_run_font(r, size=9, color=WHITE, bold=True)
    r = c2.paragraphs[0].add_run(value)
    set_run_font(r, size=12, color=NAVY, bold=True)

doc.add_paragraph()
add_anchor(doc, "Scikit-learn entrena. MLflow registra. El Consejo compara y decide.")

p = doc.add_paragraph(style="Heading 2")
p.add_run("Cómo usar este documento")
add_bullet(doc, "No intentes memorizarlo palabra por palabra. **Lee primero cada frase ancla** y cuenta el resto con tus palabras.")
add_bullet(doc, "Las cajas azules indican **qué señalar, cuándo cambiar de diapositiva o cuándo entrar en la app**.")
add_bullet(doc, "El texto principal está escrito para sonar hablado: incluye pausas, ejemplos y frases de transición.")
add_bullet(doc, "Duración objetivo: **unos 17 minutos y medio**, con margen para respirar o responder una interrupción.")

doc.add_page_break()

p = doc.add_paragraph(style="Heading 1")
p.add_run("Antes de empezar · preparación mínima")
add_bullet(doc, "Deja la presentación abierta en la diapositiva 1 y la app abierta en otra pestaña.")
add_bullet(doc, "Si vas a enseñar la app en directo, lleva preparada una **sesión de prueba**, un equipo ya dentro y, si es posible, tres runs completadas. Así no dependes del tiempo de entrenamiento.")
add_bullet(doc, "Ten localizada una matriz de confusión y el informe JSON para que el comparador no se quede en una tabla de decimales.")
add_bullet(doc, "Si falla Internet, continúa con las diapositivas 6 a 10: el guion no depende de que la demo funcione.")

add_section_title(doc, "1", "Apertura", "0:45")
add_cue(doc, "Mira al público antes de hablar. No leas el subtítulo completo; úsalo solo como contexto.")
add_rich_paragraph(doc, "Hoy os voy a presentar **La Cámara de los Experimentos Perdidos**, un proyecto para aprender MLflow de una forma bastante práctica y, sobre todo, sin separar la parte técnica de la toma de decisiones.")
add_rich_paragraph(doc, "La idea utiliza el dataset Digits de scikit-learn, modelos Random Forest y una aplicación en Streamlit. Toda la estética está inspirada en la Biblioteca Jedi, pero la narrativa no está puesta solo para decorar: la usamos como una especie de mapa mental para entender qué guarda MLflow y por qué lo guarda.")
add_rich_paragraph(doc, "Voy a dividir la explicación en dos partes. Primero quiero dejar claro **qué problema resuelve MLflow y cómo funciona el tracking**. Después os enseño la aplicación y cómo convierte esos conceptos en una actividad de equipo.")
add_anchor(doc, "Primero entendemos la memoria del experimento; después vemos cómo la practicamos en la app.")

add_section_title(doc, "2", "El problema: entrenar no es suficiente", "1:15")
add_cue(doc, "Señala los tres nombres de archivo y deja una pausa breve para que se reconozca la situación.")
add_rich_paragraph(doc, "Quería empezar por una escena que creo que nos ha pasado a casi todos: **modelo_final, modelo_final_bueno, modelo_final_ahora_si**. Al principio parece una tontería, pero en cuanto hacemos varios experimentos aparece un problema serio.")
add_rich_paragraph(doc, "Podemos tener un modelo guardado y, aun así, no saber con qué datos se entrenó, qué hiperparámetros tenía, qué métricas obtuvo, qué clases confundía o por qué decidimos quedarnos con él. Es decir, conservamos el resultado final, pero perdemos la historia que permite entenderlo y repetirlo.")
add_rich_paragraph(doc, "Y para mí esa es la idea de partida del proyecto: **experimentar mucho no es el problema**. De hecho, queremos probar alternativas. El problema aparece cuando cada prueba queda aislada, sin contexto y sin evidencias comparables.")
add_anchor(doc, "El problema no es experimentar muchas veces; el problema es perder la historia de cada experimento.")

add_section_title(doc, "3", "Qué es MLflow Tracking", "2:00")
add_cue(doc, "Recorre visualmente los cuatro bloques: decisiones previas, resultados, archivos y reproducibilidad.")
add_rich_paragraph(doc, "MLflow Tracking es, básicamente, **un sistema de registro y organización de ejecuciones de machine learning**. No sustituye a scikit-learn, no entrena por nosotros y tampoco elige el mejor modelo. Su trabajo es conservar de forma ordenada lo que ocurrió en cada ejecución.")
add_rich_paragraph(doc, "Antes del entrenamiento tenemos decisiones previas: por ejemplo, usar 50, 100 o 300 árboles, limitar o no la profundidad, fijar el tamaño del conjunto de prueba y guardar el random_state. Eso son los **parameters**, porque describen la configuración con la que entramos al entrenamiento.")
add_rich_paragraph(doc, "Después de entrenar obtenemos resultados numéricos: accuracy, precision, recall o F1. Eso son las **metrics**. La diferencia parece sencilla, pero es esencial: un parámetro expresa una decisión tomada antes; una métrica expresa un resultado observado después.")
add_rich_paragraph(doc, "Luego están los **artifacts**, que son archivos con más contexto que una cifra aislada. En este proyecto guardamos una matriz de confusión en PNG y un classification report en JSON. La accuracy puede decirme que el modelo acierta alrededor del 96 %, pero la matriz me permite ver dónde se equivoca y qué dígitos está confundiendo.")
add_rich_paragraph(doc, "También registramos el **modelo entrenado**. En nuestro caso se guarda mediante la integración de MLflow con scikit-learn, junto con una firma de entradas y salidas y un pequeño ejemplo de entrada. Así no conservamos solo el número final: conservamos un objeto que se puede cargar y cuyo contrato de datos queda documentado.")
add_rich_paragraph(doc, "Todo esto aporta trazabilidad. Si una run tiene un identificador, sus parámetros, métricas, artefactos, modelo y etiquetas, otra persona puede comparar lo que hicimos y acercarse mucho más a reproducirlo. Digo ‘acercarse’ porque MLflow ayuda muchísimo, pero una reproducción perfecta también exige controlar datos, código, dependencias y entorno.")
add_anchor(doc, "MLflow no entrena ni decide: conserva las evidencias necesarias para comparar y reproducir.")

add_section_title(doc, "4", "La Biblioteca del Templo Jedi", "2:00")
add_cue(doc, "Baja por la tabla de arriba abajo. No leas el código completo; explica la función de cada pieza.")
add_rich_paragraph(doc, "Aquí es donde entra la analogía Jedi. La **Biblioteca** representa MLflow Tracking completo: el lugar donde queda organizado el conocimiento. Dentro de esa biblioteca, una **campaña** equivale a un Experiment, es decir, el contenedor que agrupa ejecuciones relacionadas.")
add_rich_paragraph(doc, "Cada **misión** es una Run. Una run representa una ejecución concreta: un entrenamiento con una configuración y unos resultados determinados. Por eso, si entreno tres Random Forest distintos, necesito tres runs, aunque estén dentro del mismo Experiment.")
add_rich_paragraph(doc, "La **estrategia** son los Parameters: lo que decidimos antes de entrenar. El **informe del Consejo** son las Metrics: lo que observamos al evaluar. El **holocrón** es un Artifact: una evidencia que conserva información rica, como una imagen o un informe. Y el **Jedi entrenado** representa el Model que podemos guardar y recuperar.")
add_rich_paragraph(doc, "En el código, el orden sería muy parecido a este: primero indicamos dónde se va a guardar el tracking; después seleccionamos o creamos el Experiment; abrimos una Run con start_run; y, dentro de esa run, registramos parámetros, métricas, artefactos y modelo.")
add_rich_paragraph(doc, "En esta aplicación el tracking se guarda en una base SQLite de MLflow y los artefactos quedan en el directorio de datos. Además, cada Experiment incorpora el código de sesión y cada run lleva etiquetas de equipo y sesión. Eso permite que una actividad con varios grupos no termine mezclada en un único bloque imposible de rastrear.")
add_rich_paragraph(doc, "La última fila es intencionada: el **Consejo Jedi no es una función de MLflow**. Es el equipo humano que interpreta las evidencias. MLflow organiza la información, pero el criterio sigue siendo nuestro.")
add_anchor(doc, "Experiment agrupa; Run representa una ejecución; Parameters describen; Metrics evalúan; Artifacts explican; Model conserva.")

add_section_title(doc, "5", "El flujo de un experimento", "1:45")
add_cue(doc, "Sigue el diagrama con el dedo o el cursor: datos → entrenamiento → evaluación → registro → comparación → decisión.")
add_rich_paragraph(doc, "Si juntamos las piezas, este es el recorrido completo. Partimos del dataset Digits, preparamos los datos y entrenamos un Random Forest con scikit-learn. Después generamos predicciones y calculamos las métricas.")
add_rich_paragraph(doc, "A continuación MLflow registra la configuración y los resultados de esa ejecución. En nuestro proyecto son **cinco parámetros**, **cuatro métricas globales más dos recalls por clase**, **dos artefactos** y **un modelo**. Repetimos el proceso con tres configuraciones y obtenemos tres runs comparables.")
add_rich_paragraph(doc, "La comparación viene después. Y esto es importante: no elegimos automáticamente la run con la accuracy más alta. Miramos el contexto, el coste, el comportamiento por clase y las evidencias. La decisión final puede cambiar según el problema que estemos intentando resolver.")
add_rich_paragraph(doc, "Con esto cierro la parte conceptual. Ya sabemos qué memoria conserva MLflow. Ahora voy a enseñar cómo la aplicación obliga a utilizar esa memoria para argumentar una decisión.")
add_anchor(doc, "Los datos y scikit-learn producen el modelo; MLflow conserva su historia; las personas interpretan la evidencia.")

doc.add_page_break()

add_section_title(doc, "6", "La aplicación", "1:15")
add_cue(doc, "Si haces demo, abre ahora la pestaña de Streamlit. Si no, mantén esta diapositiva y describe el mismo recorrido.", label="CAMBIO A LA APP")
add_rich_paragraph(doc, "La aplicación se llama igual que el proyecto y está construida con Streamlit. Hay dos entradas: **Profesorado** y **Alumnado**. El profesorado crea una sesión y obtiene un código con formato JEDI-XXXX. El alumnado no crea cuentas: introduce ese código y elige uno de diez nombres de equipo disponibles.")
add_rich_paragraph(doc, "El trabajo está pensado para grupos de tres. El **Navegante** controla el recorrido, el **Escriba** introduce el código y el **Auditor** comprueba que lo que se registra corresponde realmente a parámetros, métricas o evidencias. Estos roles ayudan a que no sea una actividad de una sola persona delante del teclado.")
add_rich_paragraph(doc, "Desde el panel docente se puede ver en qué cámara está cada equipo, sus intentos, pistas, puntuación y fase. También se puede enviar un aviso global, desbloquear un progreso, reiniciar un equipo, reasignar un holocrón y exportar los resultados en JSON o CSV.")
add_anchor(doc, "La app convierte el tracking en un recorrido colaborativo, no en una lista de funciones para memorizar.")

add_section_title(doc, "7", "Las seis cámaras y el marcador", "1:45")
add_cue(doc, "En la app, entra en Mapa y luego en Cámara actual. Enseña una teoría breve, el editor y el botón de pista.")
add_rich_paragraph(doc, "Cada sesión selecciona **seis preguntas de un banco de cincuenta**. Hay una pregunta por cada concepto esencial: Experiment, Run, Parameters, Metrics, Artifacts y Model. Así cambia la actividad entre sesiones, pero nunca perdemos la cobertura del contenido.")
add_rich_paragraph(doc, "Los retos no son todos iguales. Algunos son de elección múltiple; otros piden completar código, detectar una llamada equivocada, corregirla o clasificar una información. Cada cámara incluye una explicación corta, el contexto necesario, un editor y tres pistas progresivas.")
add_rich_paragraph(doc, "La puntuación hace visible el progreso sin convertir el error en un bloqueo. Resolver una cámara suma 20 puntos. Una pista resta 5. Hay cinco oportunidades y, si el equipo falla las cinco, aparece la solución, se aplican 10 puntos de penalización y puede continuar.")
add_rich_paragraph(doc, "Esa decisión pedagógica me parece importante: equivocarse tiene una consecuencia, pero nadie se queda encerrado en la primera parte y pierde toda la práctica posterior.")
add_anchor(doc, "La variación evita memorizar respuestas; el avance asistido evita que un error detenga todo el aprendizaje.")

add_section_title(doc, "8", "Seguridad y tres misiones", "2:15")
add_cue(doc, "En Laboratorio, muestra las tarjetas de Tatooine, Coruscant y Mustafar y abre el editor de una misión.")
add_rich_paragraph(doc, "Esta parte tenía un reto técnico claro: quería que el alumnado escribiera sintaxis real de MLflow, pero no podía ejecutar directamente cualquier texto escrito en un editor web.")
add_rich_paragraph(doc, "La solución fue separar **validación** y **ejecución**. El código del alumnado se convierte en un árbol sintáctico, un AST. El validador revisa qué llamadas aparecen, si hay imports, si se usa un bloque with cuando corresponde, si se están mezclando parámetros y métricas y si están presentes todas las piezas exigidas.")
add_rich_paragraph(doc, "Si la estructura es válida, la aplicación **ignora ese texto para la ejecución** y lanza una plantilla interna controlada. Esa plantilla carga Digits, hace una división estratificada con un 20 % de test y random_state 42, entrena el Random Forest, calcula las métricas, genera los dos artefactos y registra una run real en MLflow. Así conservamos el valor de escribir código sin abrir ejecución arbitraria en el servidor.")
add_rich_paragraph(doc, "Las misiones aumentan la autonomía. **Tatooine** está guiada y usa 50 árboles con profundidad máxima 8. **Coruscant** tiene 100 árboles, profundidad 10 y un sabotaje: parámetros y métricas están intercambiados y hay que corregirlos. **Mustafar** pide el tracking completo casi desde cero, con 300 árboles y sin límite de profundidad.")
add_rich_paragraph(doc, "El recordatorio 5-4-2-1 resume el mínimo: cinco parámetros, cuatro métricas globales, dos artefactos y un modelo. Además, la ejecución real añade recall del dígito 1 y del dígito 8 como métricas diagnósticas, porque luego serán útiles para tomar decisiones que no dependan solo de la accuracy.")
add_anchor(doc, "El alumnado escribe y razona el tracking; la aplicación valida la estructura; solo una plantilla interna segura entrena y registra.")

add_section_title(doc, "9", "Comparador de runs y Orden 66", "2:15")
add_cue(doc, "En Comparador, señala primero parámetros y métricas; después abre una matriz o el JSON. Pulsa Orden 66 solo si la sesión de demo está preparada.")
add_rich_paragraph(doc, "Cuando terminan las tres misiones, ya no tenemos un único modelo: tenemos tres runs con distinta complejidad. Los resultados orientativos muestran que Mustafar obtiene la mejor accuracy y F1 global, y también mejora mucho el recall del dígito 1. Pero el recall del 8 se mantiene alrededor de 0,8571 en las tres.")
add_rich_paragraph(doc, "Ese dato sirve para demostrar que **una cifra global puede ocultar un comportamiento concreto**. Si el problema exige reconocer ochos, Mustafar no resuelve automáticamente la limitación aunque tenga más árboles y una accuracy mayor. Por eso el comparador permite abrir la matriz de confusión y el classification report: necesitamos ver el patrón de error, no solo el promedio.")
add_rich_paragraph(doc, "Después aparece ‘Ejecuten Orden 66’. No cambia las runs ni fabrica resultados nuevos. Asigna al equipo uno de siete criterios de decisión: por ejemplo, recursos limitados, auditoría, comunicación a un público no técnico, recall del 1, recall del 8 o preparación para producción.")
add_rich_paragraph(doc, "Y aquí está el objetivo final del proyecto: con los mismos datos, dos equipos pueden recomendar runs distintas si su criterio y su evidencia están bien explicados. El equipo con recursos limitados puede preferir una mejora más modesta y barata; el equipo centrado en rescatar unos priorizará recall_1; el de auditoría mirará parámetros, artefactos, modelo y run_id.")
add_rich_paragraph(doc, "El holocrón cambia la pregunta que hacemos a los resultados. **No cambia los resultados**. Esa diferencia obliga a argumentar y evita que la actividad termine en ‘gana el decimal más alto’.")
add_anchor(doc, "La accuracy resume; los artefactos explican; el criterio de negocio decide qué evidencia importa.")

doc.add_page_break()
add_section_title(doc, "10", "Defensa ante el Consejo y cierre", "1:30")
add_cue(doc, "Vuelve a la presentación. Haz una pausa antes de la frase final grande.", label="REGRESO A DIAPOSITIVAS")
add_rich_paragraph(doc, "El recorrido termina con una defensa de 45 a 60 segundos. Cada equipo tiene que elegir una run, compararla con otra, citar al menos dos métricas, indicar qué artefacto consultó, explicar un hallazgo, reconocer una limitación y terminar con una recomendación.")
add_rich_paragraph(doc, "La plantilla les da estructura, pero no les da la respuesta: ‘Elegimos esta run porque… Consultamos este artefacto y observamos… Descartamos esta otra por… Nuestra decisión tiene esta limitación… y recomendamos…’. Cuando la sellan, el profesorado puede revisarla desde su panel.")
add_rich_paragraph(doc, "Si tuviera que resumir todo el proyecto en una sola idea, sería esta: **un buen experimento no es solo un modelo con una métrica alta; es una decisión que otra persona puede entender, revisar y discutir con evidencias**.")
add_rich_paragraph(doc, "La Biblioteca del Templo Jedi impedía que el conocimiento se perdiera. MLflow hace lo mismo con nuestros experimentos. Scikit-learn entrena. MLflow registra. Y el Consejo, es decir, nosotros, compara y decide.")
add_anchor(doc, "Un experimento útil no termina al entrenar: termina cuando podemos explicar y defender la decisión.")

doc.add_page_break()
p = doc.add_paragraph(style="Heading 1")
p.add_run("Control del tiempo")
p = doc.add_paragraph(style="Heading 2")
p.add_run("Si necesitas dejarlo en 15 minutos")
add_bullet(doc, "Diapositiva 3: resume Parameters, Metrics, Artifacts y Model en una sola frase cada uno.")
add_bullet(doc, "Diapositiva 4: explica la analogía sin leer las llamadas de código.")
add_bullet(doc, "Diapositiva 7: menciona el banco de 50, las seis categorías y el desbloqueo tras cinco fallos; omite el detalle completo de puntuación.")
add_bullet(doc, "Diapositiva 9: abre solo un artefacto y explica un único holocrón.")

p = doc.add_paragraph(style="Heading 2")
p.add_run("Si puedes acercarte a 20 minutos")
add_bullet(doc, "En la diapositiva 3, añade un ejemplo completo: n_estimators=100 como parámetro, accuracy=0,9611 como métrica y matriz_confusion.png como artefacto.")
add_bullet(doc, "En la diapositiva 6, enseña brevemente el panel del profesorado y el código JEDI-XXXX.")
add_bullet(doc, "En la diapositiva 8, muestra el sabotaje de Coruscant y pregunta al público qué llamada está mal.")
add_bullet(doc, "En la diapositiva 9, compara Mustafar y Coruscant y abre la fila del dígito 8 en el informe JSON.")

p = doc.add_paragraph(style="Heading 1")
p.add_run("Plan B si la app no carga")
add_rich_paragraph(doc, "Puedes decir: **‘La aplicación está desplegada, pero para no depender de la conexión voy a enseñaros el recorrido desde las diapositivas. Lo importante es el flujo, porque la ejecución y el registro son exactamente los mismos.’** Después continúa con las diapositivas 6 a 10 sin disculparte más.")

doc.add_page_break()
p = doc.add_paragraph(style="Heading 1")
p.add_run("Posibles preguntas y respuestas cortas")
qa = [
    ("¿MLflow entrena el modelo?", "No. Aquí entrena scikit-learn. MLflow registra la ejecución y conserva sus evidencias."),
    ("¿Por qué SQLite?", "Porque para una actividad autocontenida permite guardar los metadatos de tracking sin desplegar un servidor adicional. Para escalar o tener varios nodos usaría un backend externo."),
    ("¿El código del alumnado se ejecuta?", "No. Se analiza con AST; si es válido, se ejecuta una plantilla interna conocida."),
    ("¿Por qué no elegir siempre la mayor accuracy?", "Porque puede ocultar errores por clase y no incorpora coste, trazabilidad ni criterio de uso."),
    ("¿Es totalmente reproducible?", "MLflow mejora mucho la reproducibilidad, pero también hay que versionar datos, código, dependencias y entorno."),
    ("¿Qué se guarda exactamente por run?", "Cinco parámetros, seis métricas en total, dos artefactos, el modelo, su firma, un ejemplo de entrada y etiquetas de equipo y sesión."),
]
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(table, [3000, 6360], indent_dxa=120)
table.style = "Table Grid"
for idx, title in enumerate(("Pregunta", "Respuesta breve")):
    cell = table.cell(0, idx)
    shade_cell(cell, "152238")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, size=10, color=WHITE, bold=True)
for question, answer in qa:
    cells = table.add_row().cells
    for idx, text in enumerate((question, answer)):
        p = cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        set_run_font(r, size=9.5, color=BLACK, bold=(idx == 0))
        set_cell_margins(cells[idx])
set_table_geometry(table, [3000, 6360], indent_dxa=120)

doc.add_page_break()
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("TARJETA DE RESCATE")
set_run_font(r, size=11, color=GOLD, bold=True)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
r = p.add_run("Diez ideas para hablar sin papel")
set_run_font(r, size=24, color=NAVY, bold=True)

anchors = [
    "1. Problema: perdemos la historia de cada experimento.",
    "2. MLflow registra y organiza; no entrena ni decide.",
    "3. Experiment agrupa runs relacionadas.",
    "4. Parameters = decisiones previas; Metrics = resultados posteriores.",
    "5. Artifacts explican los errores; Model conserva el objeto entrenado.",
    "6. La app trabaja con seis cámaras y tres misiones progresivas.",
    "7. El código se valida con AST y nunca se ejecuta directamente.",
    "8. Se registran 5 parámetros, 4 métricas globales + 2 recalls, 2 artefactos y 1 modelo.",
    "9. Orden 66 cambia el criterio de decisión, no los resultados.",
    "10. Scikit-learn entrena; MLflow registra; el Consejo compara y decide.",
]
for item in anchors:
    add_number(doc, item.split(". ", 1)[1])

add_anchor(doc, "Si te quedas en blanco, vuelve al flujo: entrenar → registrar → comparar → decidir.")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
